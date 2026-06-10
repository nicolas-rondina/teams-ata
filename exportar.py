"""Exportação da ata para Word (.docx) e PDF (.pdf).

A ata vem em Markdown (títulos com ##, listas, tabelas com |). Estas funções
convertem esse conteúdo para documentos formatados e retornam os bytes do
arquivo, prontos para salvar em disco ou oferecer como download.
"""
import io
import re
from datetime import datetime

from docx import Document
from fpdf import FPDF

# Ícones (emojis) não existem nas fontes padrão do PDF; trocamos por texto.
SUBSTITUICOES_PDF = {
    "✅": "[OK]",
    "🔄": "[em andamento]",
    "⚠️": "[atrasado]",
    "❌": "[bloqueado]",
    "📌": "-",
    "—": "-",
    "–": "-",
    "“": '"',
    "”": '"',
    "’": "'",
    "‘": "'",
}


def _limpar_inline(texto):
    """Remove marcações de ênfase do Markdown (**negrito**, *itálico*)."""
    texto = re.sub(r"\*\*(.+?)\*\*", r"\1", texto)
    texto = re.sub(r"(?<!\*)\*(?!\*)(.+?)\*(?!\*)", r"\1", texto)
    return texto


def _linha_de_tabela(linha):
    cru = linha.strip()
    return cru.startswith("|") and cru.endswith("|")


def _eh_separador(linha):
    return bool(re.match(r"^\s*\|?[:\-\s|]+\|?\s*$", linha)) and "-" in linha


def _celulas(linha):
    return [p.strip() for p in linha.strip().strip("|").split("|")]


def _metadados(nome_reuniao, tipo_reuniao, participantes):
    return [
        ("Reunião", nome_reuniao),
        ("Tipo", tipo_reuniao),
        ("Data", datetime.now().strftime("%d/%m/%Y às %H:%M")),
        ("Participantes", participantes),
    ]


# --------------------------------------------------------------------------- #
# Word (.docx)
# --------------------------------------------------------------------------- #
def gerar_docx(conteudo_ia, nome_reuniao, participantes, transcricao, tipo_reuniao="Padrão"):
    doc = Document()
    doc.add_heading("Ata de Reunião", level=0)

    meta = doc.add_paragraph()
    for rotulo, valor in _metadados(nome_reuniao, tipo_reuniao, participantes):
        run = meta.add_run(f"{rotulo}: ")
        run.bold = True
        meta.add_run(f"{valor}\n")

    _markdown_para_docx(doc, conteudo_ia)

    doc.add_page_break()
    doc.add_heading("Transcrição Completa", level=1)
    for linha in transcricao.split("\n"):
        doc.add_paragraph(linha)

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def _markdown_para_docx(doc, texto):
    linhas = texto.split("\n")
    i, n = 0, len(linhas)
    while i < n:
        linha = linhas[i]
        cru = linha.strip()
        if not cru:
            i += 1
            continue
        if cru.startswith("### "):
            doc.add_heading(_limpar_inline(cru[4:]), level=3)
        elif cru.startswith("## "):
            doc.add_heading(_limpar_inline(cru[3:]), level=2)
        elif cru.startswith("# "):
            doc.add_heading(_limpar_inline(cru[2:]), level=1)
        elif _linha_de_tabela(linha) and i + 1 < n and _eh_separador(linhas[i + 1]):
            bloco = []
            while i < n and _linha_de_tabela(linhas[i]):
                bloco.append(linhas[i])
                i += 1
            _tabela_docx(doc, bloco)
            continue
        elif cru.startswith("- ") or cru.startswith("* "):
            doc.add_paragraph(_limpar_inline(cru[2:]), style="List Bullet")
        elif cru.startswith("✅"):
            doc.add_paragraph(_limpar_inline(cru), style="List Bullet")
        else:
            doc.add_paragraph(_limpar_inline(cru))
        i += 1


def _tabela_docx(doc, bloco):
    cabecalho = _celulas(bloco[0])
    dados = [_celulas(l) for l in bloco[2:]]
    cols = len(cabecalho)

    tabela = doc.add_table(rows=1, cols=cols)
    tabela.style = "Table Grid"

    celulas_cab = tabela.rows[0].cells
    for j, texto in enumerate(cabecalho):
        celulas_cab[j].text = _limpar_inline(texto)
        for paragrafo in celulas_cab[j].paragraphs:
            for run in paragrafo.runs:
                run.bold = True

    for linha in dados:
        celulas = tabela.add_row().cells
        for j in range(cols):
            celulas[j].text = _limpar_inline(linha[j]) if j < len(linha) else ""


# --------------------------------------------------------------------------- #
# PDF (.pdf)
# --------------------------------------------------------------------------- #
class _PDF(FPDF):
    def footer(self):
        self.set_y(-15)
        self.set_font("Helvetica", "I", 8)
        self.cell(0, 10, f"Página {self.page_no()}", align="C")


def _sanitizar_pdf(texto):
    for de, para in SUBSTITUICOES_PDF.items():
        texto = texto.replace(de, para)
    return texto.encode("latin-1", "replace").decode("latin-1")


def _escreve_pdf(pdf, altura, texto):
    """Escreve um parágrafo sempre a partir da margem esquerda, com largura total."""
    pdf.set_x(pdf.l_margin)
    pdf.multi_cell(pdf.epw, altura, texto)


def gerar_pdf(conteudo_ia, nome_reuniao, participantes, transcricao, tipo_reuniao="Padrão"):
    pdf = _PDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()

    pdf.set_font("Helvetica", "B", 18)
    _escreve_pdf(pdf, 9, _sanitizar_pdf("Ata de Reunião"))
    pdf.ln(2)

    pdf.set_font("Helvetica", "", 11)
    for rotulo, valor in _metadados(nome_reuniao, tipo_reuniao, participantes):
        _escreve_pdf(pdf, 6, _sanitizar_pdf(f"{rotulo}: {valor}"))
    pdf.ln(3)

    _markdown_para_pdf(pdf, conteudo_ia)

    pdf.add_page()
    pdf.set_font("Helvetica", "B", 14)
    _escreve_pdf(pdf, 8, _sanitizar_pdf("Transcrição Completa"))
    pdf.ln(2)
    pdf.set_font("Helvetica", "", 9)
    for linha in transcricao.split("\n"):
        _escreve_pdf(pdf, 5, _sanitizar_pdf(linha))

    return bytes(pdf.output())


def _markdown_para_pdf(pdf, texto):
    linhas = texto.split("\n")
    i, n = 0, len(linhas)
    while i < n:
        linha = linhas[i]
        cru = linha.strip()
        if not cru:
            pdf.ln(2)
            i += 1
            continue
        if cru.startswith("## ") or cru.startswith("### "):
            pdf.ln(1)
            pdf.set_font("Helvetica", "B", 13)
            _escreve_pdf(pdf, 7, _sanitizar_pdf(_limpar_inline(cru.lstrip("# "))))
        elif cru.startswith("# "):
            pdf.set_font("Helvetica", "B", 15)
            _escreve_pdf(pdf, 8, _sanitizar_pdf(_limpar_inline(cru[2:])))
        elif _linha_de_tabela(linha) and i + 1 < n and _eh_separador(linhas[i + 1]):
            bloco = []
            while i < n and _linha_de_tabela(linhas[i]):
                bloco.append(linhas[i])
                i += 1
            _tabela_pdf(pdf, bloco)
            continue
        elif cru.startswith("- ") or cru.startswith("* "):
            pdf.set_font("Helvetica", "", 11)
            _escreve_pdf(pdf, 6, _sanitizar_pdf("  - " + _limpar_inline(cru[2:])))
        else:
            pdf.set_font("Helvetica", "", 11)
            _escreve_pdf(pdf, 6, _sanitizar_pdf(_limpar_inline(cru)))
        i += 1


def _tabela_pdf(pdf, bloco):
    cabecalho = _celulas(bloco[0])
    dados = [_celulas(l) for l in bloco[2:]]
    cols = len(cabecalho)

    pdf.set_font("Helvetica", "", 9)
    pdf.set_x(pdf.l_margin)
    with pdf.table(width=pdf.epw) as tabela:
        cab = tabela.row()
        for texto in cabecalho:
            cab.cell(_sanitizar_pdf(_limpar_inline(texto)))
        for linha in dados:
            registro = tabela.row()
            for j in range(cols):
                valor = linha[j] if j < len(linha) else ""
                registro.cell(_sanitizar_pdf(_limpar_inline(valor)))
    pdf.ln(2)
